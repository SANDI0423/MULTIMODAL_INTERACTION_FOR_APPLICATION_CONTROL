import aspose.words as aw
import speech_recognition as sr
import pyttsx3
listener = sr.Recognizer()
engine = pyttsx3.init()

lines=[]
def talk(text):
    engine.say(text)
    engine.runAndWait()
doc = aw.Document()
builder = aw.DocumentBuilder(doc)

def get_first():
    talk("Hi I am word...How may I help You")
    info=""
    try:
        with sr.Microphone() as source:
            print('listening....')
            voice = listener.listen(source)
            info = listener.recognize_google(voice)
            print(info)
    except:
        pass
    if info=="read":
        import docx
        talk("What is the file you want to read")
        with sr.Microphone() as source:
            print('listening....')
            voice = listener.listen(source)
            info = listener.recognize_google(voice)
        # Open the .docx file
        doc = docx.Document("D:\\"+info+".docx")
        h = []
        # Iterate through paragraphs and extract text
        for paragraph in doc.paragraphs:
            text = paragraph.text
            h.append(text)

        talk(" ".join(h[1:len(h)-1]))
    if info=="open":
        talk("What content do you want to type            listening")
        get_info()

def get_info():
    try:
        with sr.Microphone() as source:
            print('listening....')
            voice = listener.listen(source)
            info = listener.recognize_google(voice)
            print(info)
            talk(info)
            lines.append(info)
            if 'stop' in info:
                writedocx()
            else:
                get_info()
    except:
        pass
def writedocx():
    for i in lines:
        builder.writeln(i + "\n")
    talk("what is the filename you want to save")
    with sr.Microphone() as source:
        print('listening....')
        voice = listener.listen(source)
        info = listener.recognize_google(voice)
    doc.save("D:\\"+info+".docx")
    window1()
    talk("Saved your file")
    talk("Bye")


def window1():

    import tkinter as tk

    # Create the main application window
    root = tk.Tk()
    root.title("MS WORD")
    root.geometry("250x350")
    paragraph_label = tk.Label(root, text=" ".join(lines), wraplength=400, justify="left")
    paragraph_label.pack(padx=20, pady=20)

    # Run the main application loop
    root.mainloop()
# save document
get_first()